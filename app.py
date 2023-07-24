import os
import zipfile
from flask import Flask, request, send_file

from docx import Document
from openpyxl import load_workbook,Workbook
from openpyxl.styles import PatternFill

app = Flask(__name__)

def generate_txt(data):
    return f"Se remite a Sr(a) {data['nombre']} la disposición de presentarse en la entidad {data['entidad']} con su vehículo de placa {data['placa']}."

def generate_docx(data):
    doc = Document()
    text = generate_txt(data)
    doc.add_paragraph(text)
    return doc

def generate_xlsx(data):
    wb = Workbook()
    ws = wb.active
    text = generate_txt(data)
    ws['A1'] = text
    return wb

@app.route('/generate_document', methods=['POST'])
def generate_document():
    if request.is_json:
        data = request.get_json() 

         # Generar el archivo de texto plano
        txt_content = generate_txt(data)


        # Generar el documento DOCX
        docx_doc = generate_docx(data)
        docx_filename = 'documento.docx'
        docx_doc.save(docx_filename)

        # Generar el archivo XLSX
        xlsx_wb = generate_xlsx(data)
        xlsx_filename = 'documento.xlsx'
        xlsx_wb.save(xlsx_filename)

        # Crear un archivo ZIP para contener todos los documentos generados
        zip_filename = 'documentos.zip'
        with zipfile.ZipFile(zip_filename, 'w') as zip_file:
            zip_file.write(docx_filename)
            zip_file.write(xlsx_filename)


        # Eliminar los archivos generados
        os.remove(docx_filename)
        os.remove(xlsx_filename)


        # Devolver el archivo ZIP como adjunto en la respuesta
        return send_file(zip_filename, as_attachment=True)
    else:
        file = request.files['file']
        file_ext = os.path.splitext(file.filename)[1].lower()

        # Guardar el archivo en una ubicación temporal
        temp_filename = 'temp_file' + file_ext
        file.save(temp_filename)

        # Generar documentos según el tipo de archivo recibido
        data = {}
        if file_ext == '.xlsx':
            wb = load_workbook(temp_filename)
            ws = wb.active
            data['nombre'] = ws['A1'].value
            data['entidad'] = ws['A2'].value
            data['placa'] = ws['A3'].value
            wb.close()  # Cerrar el libro de trabajo después de leer los datos
            os.remove(temp_filename)
        elif file_ext == '.txt':
            with open(temp_filename, 'r', encoding='utf-8') as txt_file:
                lines = txt_file.readlines()
                data['nombre'] = lines[0].strip()
                data['entidad'] = lines[1].strip()
                data['placa'] = lines[2].strip()
        elif file_ext == '.docx':
            doc = Document(temp_filename)
            paragraphs = doc.paragraphs
            data['nombre'] = paragraphs[0].text.strip()
            data['entidad'] = paragraphs[1].text.strip()
            data['placa'] = paragraphs[2].text.strip()
            os.remove(temp_filename)
        else:
            os.remove(temp_filename)
            return 'Tipo de archivo no soportado. Los formatos permitidos son XLSX, TXT y DOCX.', 400

        # Generar el archivo de texto plano
        txt_content = generate_txt(data)
       

        # Generar el documento DOCX
        docx_doc = generate_docx(data)
        docx_filename = 'documento.docx'
        docx_doc.save(docx_filename)

        # Generar el archivo XLSX
        xlsx_wb = generate_xlsx(data)
        xlsx_filename = 'documento.xlsx'
        xlsx_wb.save(xlsx_filename)

        # Crear un archivo ZIP para contener todos los documentos generados
        zip_filename = 'documentos.zip'
        with zipfile.ZipFile(zip_filename, 'w') as zip_file:
            zip_file.write(docx_filename)
            zip_file.write(xlsx_filename)


        # Eliminar los archivos generados
        os.remove(docx_filename)
        os.remove(xlsx_filename)


        # Devolver el archivo ZIP como adjunto en la respuesta
        return send_file(zip_filename, as_attachment=True)

if __name__ == '__main__':
    app.run(debug=True)
